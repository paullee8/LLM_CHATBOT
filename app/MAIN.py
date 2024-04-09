import streamlit as st
from streamlit_chat import message
import os, re, copy, time, base64, asyncio
os.environ["TOKENIZERS_PARALLELISM"] = "false"
import pandas as pd
import datetime
import docx

from langchain_openai.chat_models import AzureChatOpenAI
from langchain.prompts import (ChatPromptTemplate,PromptTemplate,SystemMessagePromptTemplate,HumanMessagePromptTemplate,PipelinePromptTemplate)
from langchain.chains import LLMChain
from langchain_core.output_parsers import StrOutputParser
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.memory import ConversationBufferMemory, VectorStoreRetrieverMemory
import faiss
from langchain_community.docstore import InMemoryDocstore
from langchain_community.vectorstores import FAISS, Chroma
from langchain.schema import Document
from langchain_core.runnables import RunnablePassthrough

#UDF
from bot_mode import bot_mode
from llm_model import model_seletion
from main_llm_result import main_llm_result
from function.search_docu import search_docu
from function.Retriever_pdf_txt import retrieve_docu_qa
from function.Retriever_url import retrieve_url_qa
from function.pandas_agent import pandas_agent_module

from dotenv import load_dotenv
load_dotenv()
openai_api_base_value = os.environ.get('openai_api_base_value')
openai_api_key_value = os.environ.get('openai_api_key_value')
openai_api_version_value = os.environ.get('openai_api_version_value')
deployment_name_value = os.environ.get('deployment_name_value')
openai_api_type_value = os.environ.get('openai_api_type_value')

HuggingFaceEmbeddings_model_name = 'infgrad/stella-base-en-v2'

if 'docu_name' not in st.session_state:
    st.session_state['docu_name'] = []
if 'docu_display' not in st.session_state:
    st.session_state['docu_display'] = []
if 'pdf_path' not in st.session_state:
    st.session_state['pdf_path'] = []
if 'docu_path_df' not in st.session_state:
    st.session_state['docu_path_df'] = []

def clear_text():
    st.session_state['generated'] = []
    st.session_state['past'] = []
    st.session_state['history'] = ''

def display_docu_step_1(file):
    base64_pdf = base64.b64encode(file.read()).decode('utf-8')
    pdf_display = F'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf"></iframe>'
    return pdf_display 

def display_docu_step_2(docu):
    st.markdown(docu, unsafe_allow_html=True)

async def embedding_pdf_path(file,file_name): 
    file_path = os.path.join(os.getcwd(), file.name)
    file_path = file_path.split('/')
    file_path.insert(-1,'temp_docu_storage')
    file_path = "/".join(file_path)
    with open(file_path, "wb") as f:
        f.write(file.getvalue())
    pdf_info = Document(page_content=str(file_name), metadata={'path':file_path})
    st.session_state.pdf_path.append(pdf_info)
    db = FAISS.from_documents(st.session_state.pdf_path, HuggingFaceEmbeddings())
    st.session_state.docu_path_df = db
    st.session_state.doc_yn = True
    return db

async def main():
    global llm_model
    with st.sidebar:
        st.sidebar.button("Clear Text", on_click=clear_text)

        lang_mode = st.radio(
            ":earth_asia:",
            ("English", "Traditional Chinese", "Other")
        )
        if lang_mode == 'Other':
            lang_mode = st.sidebar.text_input("Enter your preference")

        bot_rule_input = st.sidebar.selectbox(
            "Type of chatbot",
            ("IT dog", "General", "Business")
        )
        prompt_mode = st.radio(
            "Choose a Prompt style",
            ("Standard", "Customize") 
        )
        if prompt_mode == 'Customize':
            customized_prompt_input = st.sidebar.text_input("Enter your customized prompt")
        else:
            customized_prompt_input = ''

        temperature_levels = st.slider(
            "Temperature level",  0.0, 1.0, (0.1), 0.01
            )

        url_input = st.sidebar.text_input("URL")
        Search_with_source =False
        if len(url_input)>0:
            switch_state = st.toggle("Search with source")
            if switch_state:
                Search_with_source =True
            else:
                Search_with_source =False

        df_dict = {}
        temp_docu_lt = []
        uploaded_files = st.file_uploader("Upload a file", type=None , accept_multiple_files=True) 

        if len(uploaded_files) >=1:
            st.session_state.file_yn = True
            for file in uploaded_files:
                file_name = re.sub(r'\W+', '_', file.name)
                if 'csv' in file.type:
                    file_name = file_name +'_df'
                    df = pd.read_csv(file)
                    df_dict[file_name] = df
                    st.session_state.df_yn = True
                if '.xlsx' in file.name:
                    file_name = file_name +'_df'
                    df = pd.read_excel(file)
                    df_dict[file_name] = df
                    st.session_state.df_yn = True

                if 'pdf' in file.type or 'plain' in file.type:
                    st.session_state.doc_yn = True
                    temp_docu_lt.append(file_name)
                    if file_name not in st.session_state.docu_name:
                        file_document = display_docu_step_1(file)
                        st.session_state.docu_name.append(file_name)
                        st.session_state.docu_display.append(file_document)
                        embedd_db_result = asyncio.create_task(embedding_pdf_path(file,file_name))
                
        else:
            st.session_state.df_yn = False
            st.session_state.file_yn = False
            st.session_state.doc_yn = False

    temp_docu_lt = list(set(temp_docu_lt))
    st.session_state.docu_name = list(set(st.session_state.docu_name) & set(temp_docu_lt))
    df = pd.DataFrame(st.session_state.docu_name , columns=['File stored'])
    st.sidebar.dataframe(df,hide_index=True,width=500)

    df_options = st.sidebar.multiselect(
                                    'DataFrame',
                                    df_dict)
    
    if st.sidebar.button("Save Conversation"):
        chat_name = datetime.datetime.now()
        chat_name = chat_name.strftime('%Y_%m_%d_%I_%M_%p') + "_chat_history"
        file_path = os.path.join(os.getcwd(), chat_name)
        file_path = file_path.split('/')
        file_path.insert(-1,'temp_docu_storage')
        file_path = "/".join(file_path)
        doc = docx.Document()
        doc.add_paragraph(st.session_state['history'])
        doc.save(file_path+".docx")

   
    
    llm_model, temperature_levels = model_seletion(openai_api_base_value, openai_api_version_value,deployment_name_value, 
                                                   openai_api_key_value, openai_api_type_value,temperature_levels)
    #======================================================================================================

    if st.session_state.file_yn:
        with tabs[1]:
            if len(df_dict)>0:
                for df_key,df_value in df_dict.items():
                    st.dataframe(df_value)
            if len(st.session_state.docu_name)>0:
                for file in st.session_state.docu_display:
                    display_docu_step_2(file)

    user_input = st.chat_input("Say something")
    with tabs[0]:
        bot_rule = bot_mode(bot_rule_input)
        if bot_rule == 'Experienced manager':
            st.title(f'Bot mode: :green[{bot_rule}]' )
        elif bot_rule == 'Programmer and data scientist':
            st.title('Bot mode: :grey[IT :dog:]' )
        elif bot_rule == 'Business analyst':
            st.title(f'Bot mode: :rainbow[{bot_rule}]' )

        message('Hello! How can I assist you today?', is_user=False)   

        if user_input:
            retriever_summary = None
            df_summary = None
            URL_summary = None

            if len(df_options)>0:
                try:
                    df_llm_input = pandas_agent_module(llm_model,user_input,lang_mode, df_options, df_dict)
                    df_summary = df_llm_input.pandas_agent_job()
                except Exception as e: 
                    print('error:', str(e))

            if st.session_state.doc_yn:
                try:
                    user_query_check = search_docu(llm_model,user_input, st.session_state.docu_name, st.session_state.docu_path_df)
                    if user_query_check.calculate_similarity()[1]<=1:
                        embedded_docu = retrieve_docu_qa(llm_model, lang_mode ,user_query_check.calculate_similarity()[0]['path'], user_input,HuggingFaceEmbeddings_model_name)
                        retriever_summary = embedded_docu.retrieve_query()
                except Exception as e: 
                    print('error:', str(e))
            
            if len(url_input)>0:
                if Search_with_source:
                    URL_summary = retrieve_url_qa(llm_model,lang_mode,bot_rule,url_input,user_input,HuggingFaceEmbeddings_model_name).retriever_url()

            answer_lt = [retriever_summary,df_summary,URL_summary]
            answer_lt = [i for i in answer_lt if i is not None]

            if len(answer_lt)>0:
                print('answer_lt[0]',answer_lt)
                output = answer_lt[0]
            else:
                output = main_llm_result(llm_model,lang_mode,bot_rule, prompt_mode,customized_prompt_input,
                                        user_input, st.session_state.history, df_dict)

            st.session_state.past.append(user_input)
            st.session_state.generated.append(output)
            st.session_state.history += (f"Human:{user_input}"+"\n"+f"Chatbot:{output}"+"\n")

        if st.session_state['generated']:
            for i in range(len(st.session_state['generated'])):
                message(st.session_state['past'][i], is_user=True, key=str(i) + '_user')
                message(st.session_state["generated"][i], key=str(i))
                    

if __name__ == '__main__':
    retriever_states_lt = ['file_yn','df_yn','doc_yn']
    for states in retriever_states_lt:
        if states not in st.session_state:
            st.session_state[states] = False

    if 'generated' not in st.session_state:
        st.session_state['generated'] = []
    if 'past' not in st.session_state:
        st.session_state.past = []
    if 'history' not in st.session_state:
        st.session_state['history'] = ''

    tab_titles = ['Main', 'Source']
    tabs = st.tabs(tab_titles)

    asyncio.run(main())